import os
import pdfplumber
import docx


def extract_text_from_pdf(file_path):
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text


def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = docx.Document(file_path)

        # Extract paragraph text
        for para in doc.paragraphs:
            text += para.text + "\n"

        # Extract text from tables (skills grids, experience columns, etc.)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text += cell_text + "\n"

    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text


def extract_text_from_txt(file_path):
    text = ""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    except Exception as e:
        print(f"Error reading TXT {file_path}: {e}")
    return text


def extract_resume_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    else:
        return ""